import os
import re
import io
import math
import json
import logging
import tempfile
from collections import Counter
from typing import Dict, List, Tuple, Optional
from flask import Flask, request, jsonify, send_from_directory
from werkzeug.utils import secure_filename

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Import configuration
try:
    from config import Config
    config_warnings = Config.validate()
    for w in config_warnings:
        logger.warning(w)
except ImportError:
    logger.warning("config.py not found. Using default settings.")
    class Config:
        AI_PROVIDER = "template"
        OPENAI_API_KEY = ""
        ANTHROPIC_API_KEY = ""
        OPENAI_MODEL = "gpt-4o-mini"
        ANTHROPIC_MODEL = "claude-3-haiku-20240307"
        DEBUG = True
        PORT = 5000
        SUMMARY_SENTENCES = 10
        MAX_TOKENS_SUMMARY = 500
        MAX_TOKENS_DRAFT = 1500
        UPLOAD_FOLDER = "uploads"
        MAX_FILE_SIZE = 16 * 1024 * 1024  # 16MB

# PDF Processing Libraries
try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
    logger.info("pdfplumber loaded successfully")
except ImportError:
    PDFPLUMBER_AVAILABLE = False
    logger.warning("pdfplumber not available")

try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
    logger.info("PyPDF2 loaded successfully")
except ImportError:
    PYPDF2_AVAILABLE = False
    logger.warning("PyPDF2 not available")

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
    logger.info("python-docx loaded successfully")
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not available")

# Initialize AI clients
openai_client = None
anthropic_client = None

if Config.AI_PROVIDER == "openai" and Config.OPENAI_API_KEY:
    try:
        from openai import OpenAI
        openai_client = OpenAI(api_key=Config.OPENAI_API_KEY)
        logger.info("OpenAI client initialized successfully")
    except Exception as e:
        logger.error(f"Failed to initialize OpenAI: {e}")

if Config.AI_PROVIDER == "anthropic" and Config.ANTHROPIC_API_KEY:
    try:
        import anthropic
        anthropic_client = anthropic.Anthropic(api_key=Config.ANTHROPIC_API_KEY)
        logger.info("Anthropic client initialized successfully")
    except Exception as e:
        logger.error(f"Failed to initialize Anthropic: {e}")

# Initialize Flask application
app = Flask(__name__, static_folder="static")
app.config['MAX_CONTENT_LENGTH'] = getattr(Config, 'MAX_FILE_SIZE', 16 * 1024 * 1024)


ALLOWED_EXTENSIONS = {'pdf', 'docx', 'doc', 'txt'}


def allowed_file(filename: str) -> bool:
    """Check if the uploaded file has an allowed extension."""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS



# PDF AND DOCUMENT EXTRACTION

def extract_text_from_pdf_pdfplumber(file_stream) -> Tuple[str, Dict]:
    """
    Extract text from PDF using pdfplumber.
    
    pdfplumber is excellent for maintaining layout and extracting
    text from complex PDF structures including tables.
    
    Args:
        file_stream: File-like object containing PDF data
        
    Returns:
        Tuple of (extracted_text, metadata)
    """
    if not PDFPLUMBER_AVAILABLE:
        raise ImportError("pdfplumber is not installed")
    
    text_parts = []
    metadata = {
        "page_count": 0,
        "extraction_method": "pdfplumber",
        "has_tables": False
    }
    
    with pdfplumber.open(file_stream) as pdf:
        metadata["page_count"] = len(pdf.pages)
        
        for page_num, page in enumerate(pdf.pages, 1):
            page_text = page.extract_text()
            
            if page_text:
                text_parts.append(f"\n--- Page {page_num} ---\n")
                text_parts.append(page_text)
            
            tables = page.extract_tables()
            if tables:
                metadata["has_tables"] = True
                for table in tables:
                    if table:
                        for row in table:
                            if row:
                                row_text = " | ".join(str(cell) if cell else "" for cell in row)
                                text_parts.append(row_text)
    
    return "\n".join(text_parts), metadata


def extract_text_from_pdf_pypdf2(file_stream) -> Tuple[str, Dict]:
    """
    Extract text from PDF using PyPDF2.
    
    PyPDF2 is a fallback option that works well for simpler PDFs.
    
    Args:
        file_stream: File-like object containing PDF data
        
    Returns:
        Tuple of (extracted_text, metadata)
    """
    if not PYPDF2_AVAILABLE:
        raise ImportError("PyPDF2 is not installed")
    
    text_parts = []
    metadata = {
        "page_count": 0,
        "extraction_method": "PyPDF2"
    }
    
    # Reset stream position
    file_stream.seek(0)
    
    reader = PyPDF2.PdfReader(file_stream)
    metadata["page_count"] = len(reader.pages)
    
    # Extract document info if available
    if reader.metadata:
        if reader.metadata.title:
            metadata["title"] = reader.metadata.title
        if reader.metadata.author:
            metadata["author"] = reader.metadata.author
    
    for page_num, page in enumerate(reader.pages, 1):
        page_text = page.extract_text()
        if page_text:
            text_parts.append(f"\n--- Page {page_num} ---\n")
            text_parts.append(page_text)
    
    return "\n".join(text_parts), metadata


def extract_text_from_pdf(file_stream) -> Tuple[str, Dict]:
    """
    Extract text from PDF using the best available method.
    
    Tries pdfplumber first (better quality), falls back to PyPDF2.
    
    Args:
        file_stream: File-like object containing PDF data
        
    Returns:
        Tuple of (extracted_text, metadata)
    """
    errors = []
    
    # Try pdfplumber first (usually better quality)
    if PDFPLUMBER_AVAILABLE:
        try:
            file_stream.seek(0)
            return extract_text_from_pdf_pdfplumber(file_stream)
        except Exception as e:
            errors.append(f"pdfplumber: {str(e)}")
            logger.warning(f"pdfplumber extraction failed: {e}")
    
    # Fall back to PyPDF2
    if PYPDF2_AVAILABLE:
        try:
            file_stream.seek(0)
            return extract_text_from_pdf_pypdf2(file_stream)
        except Exception as e:
            errors.append(f"PyPDF2: {str(e)}")
            logger.warning(f"PyPDF2 extraction failed: {e}")
    
    # If both fail
    error_msg = " | ".join(errors) if errors else "No PDF library available"
    raise Exception(f"PDF extraction failed: {error_msg}")


def extract_text_from_docx(file_stream) -> Tuple[str, Dict]:
    """
    Extract text from DOCX files using python-docx.
    
    Args:
        file_stream: File-like object containing DOCX data
        
    Returns:
        Tuple of (extracted_text, metadata)
    """
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx is not installed")
    
    file_stream.seek(0)
    doc = DocxDocument(file_stream)
    
    text_parts = []
    metadata = {
        "paragraph_count": 0,
        "extraction_method": "python-docx"
    }
    
    # Extract core properties if available
    if doc.core_properties:
        if doc.core_properties.title:
            metadata["title"] = doc.core_properties.title
        if doc.core_properties.author:
            metadata["author"] = doc.core_properties.author
    
    # Extract paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)
            metadata["paragraph_count"] += 1
    
    # Extract tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                text_parts.append(row_text)
    
    return "\n\n".join(text_parts), metadata


def extract_text_from_txt(file_stream) -> Tuple[str, Dict]:
    """
    Extract text from plain text files.
    
    Args:
        file_stream: File-like object containing text data
        
    Returns:
        Tuple of (extracted_text, metadata)
    """
    file_stream.seek(0)
    
    encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
    content = None
    used_encoding = None
    
    for encoding in encodings:
        try:
            file_stream.seek(0)
            content = file_stream.read()
            if isinstance(content, bytes):
                content = content.decode(encoding)
            used_encoding = encoding
            break
        except (UnicodeDecodeError, AttributeError):
            continue
    
    if content is None:
        raise Exception("Could not decode text file with any supported encoding")
    
    metadata = {
        "extraction_method": "plain_text",
        "encoding": used_encoding
    }
    
    return content, metadata


def process_uploaded_file(file) -> Tuple[str, Dict]:
    """
    Process an uploaded file and extract text based on file type.
    
    Args:
        file: Werkzeug FileStorage object
        
    Returns:
        Tuple of (extracted_text, metadata)
    """
    filename = secure_filename(file.filename)
    file_ext = filename.rsplit('.', 1)[1].lower() if '.' in filename else ''
    
    # Read file into memory
    file_content = io.BytesIO(file.read())
    
    metadata = {
        "filename": filename,
        "file_type": file_ext
    }
    
    if file_ext == 'pdf':
        text, extract_meta = extract_text_from_pdf(file_content)
    elif file_ext in ['docx', 'doc']:
        text, extract_meta = extract_text_from_docx(file_content)
    elif file_ext == 'txt':
        text, extract_meta = extract_text_from_txt(file_content)
    else:
        raise ValueError(f"Unsupported file type: {file_ext}")
    
    metadata.update(extract_meta)
    
    # Clean the extracted text
    text = clean_extracted_text(text)
    
    metadata["character_count"] = len(text)
    metadata["word_count_raw"] = len(text.split())
    
    return text, metadata


def clean_extracted_text(text: str) -> str:
    """
    Clean text extracted from documents.
    
    Handles common PDF extraction artifacts like:
    - Excessive whitespace
    - Page markers
    - Header/footer remnants
    - Hyphenation at line breaks
    
    Args:
        text: Raw extracted text
        
    Returns:
        Cleaned text
    """
    # Remove page markers we added
    text = re.sub(r"\n--- Page \d+ ---\n", "\n\n", text)
    
    # Fix hyphenation at line breaks (word- \n continuation)
    text = re.sub(r"(\w)-\s*\n\s*(\w)", r"\1\2", text)
    
    text = re.sub(r"[ \t]+", " ", text)
    
    text = re.sub(r"\n{3,}", "\n\n", text)
    
    text = re.sub(r"\x00", "", text)  # Null characters
    text = re.sub(r"[\x01-\x08\x0b\x0c\x0e-\x1f]", "", text)  # Control characters
    
    text = re.sub(r"^\s*\d+\s*$", "", text, flags=re.MULTILINE)
    
    text = re.sub(r"Page \d+ of \d+", "", text, flags=re.IGNORECASE)
    
    return text.strip()


# NLP UTILITIES - Text Preprocessing and Analysis

STOP_WORDS = {
    "a", "an", "the", "and", "or", "but", "in", "on", "at", "to", "for", "of",
    "with", "by", "from", "up", "about", "into", "through", "during", "is",
    "are", "was", "were", "be", "been", "being", "have", "has", "had", "do",
    "does", "did", "will", "would", "could", "should", "may", "might", "shall",
    "can", "need", "ought", "this", "that", "these", "those", "it", "its",
    "they", "them", "their", "we", "our", "you", "your", "he", "she", "him",
    "her", "his", "i", "me", "my", "also", "such", "as", "which", "who",
    "when", "where", "how", "what", "all", "each", "every", "both", "more",
    "most", "other", "some", "any", "so", "then", "than", "only", "very",
    "just", "not", "no", "nor", "neither", "either", "if", "while", "although",
    "because", "since", "unless", "until", "after", "before", "whether",
    "though", "even", "under", "over", "further", "same", "between", "out",
    "off", "again", "once", "here", "there", "own", "few", "too", "per",
    "etc", "within", "without", "across", "s", "t", "re", "ve", "ll", "d",
    "m", "o", "p", "q", "r", "u", "v", "w", "x", "y", "z", "must", "based",
    "including", "following", "using", "according", "related", "upon", "page"
}


def clean_text(text: str) -> str:
    """Clean and normalize policy document text."""
    text = re.sub(
        r"National Digital Government.*?Draft V[\d.]+",
        "",
        text,
        flags=re.IGNORECASE | re.DOTALL
    )
    text = re.sub(r"https?://\S+", "", text)
    text = re.sub(r"\d+\s*$", "", text, flags=re.MULTILINE)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]+", " ", text)
    return text.strip()


def tokenize(text: str) -> List[str]:
    """Tokenize text into words."""
    return re.findall(r"\b[a-zA-Z][a-zA-Z-]*[a-zA-Z]\b", text.lower())


def sentence_split(text: str) -> List[str]:
    """Split text into sentences."""
    sentences = re.split(r"(?<=[.!?])\s+(?=[A-Z])", text)
    expanded = []
    for s in sentences:
        parts = re.split(r"(?:^|\n)\s*(?:\d+\.|•|–)\s*", s)
        expanded.extend(parts)
    return [s.strip() for s in expanded if len(s.strip()) > 25]


def compute_tf(tokens: List[str]) -> Dict[str, float]:
    """Compute Term Frequency (TF) for tokens."""
    filtered = [t for t in tokens if t not in STOP_WORDS and len(t) > 3]
    total = len(filtered) or 1
    counts = Counter(filtered)
    return {word: count / total for word, count in counts.items()}


def compute_idf(sentences: List[str]) -> Dict[str, float]:
    """Compute Inverse Document Frequency (IDF) for terms."""
    n_sentences = len(sentences) or 1
    df = Counter()
    
    for sentence in sentences:
        words = set(tokenize(sentence))
        for word in words:
            if word not in STOP_WORDS and len(word) > 3:
                df[word] += 1
    
    idf = {}
    for word, freq in df.items():
        idf[word] = math.log((n_sentences + 1) / (freq + 1)) + 1
    
    return idf


def score_sentence_tfidf(sentence: str, tf: Dict[str, float], 
                          idf: Dict[str, float]) -> float:
    """Score a sentence using TF-IDF weighting."""
    words = tokenize(sentence)
    if not words:
        return 0.0
    
    score = 0.0
    for word in words:
        if word not in STOP_WORDS:
            tf_score = tf.get(word, 0)
            idf_score = idf.get(word, 1)
            score += tf_score * idf_score
    
    return score / math.sqrt(len(words))


def extract_key_phrases(text: str, n: int = 20) -> List[str]:
    """Extract key phrases/terms from text using frequency analysis."""
    tokens = tokenize(text)
    filtered = [t for t in tokens if t not in STOP_WORDS and len(t) > 4]
    counts = Counter(filtered)
    return [word for word, _ in counts.most_common(n)]


def extractive_summarise(text: str, n_sentences: int = 10) -> List[str]:
    """Generate extractive summary using TF-IDF sentence scoring."""
    sentences = sentence_split(text)
    
    if len(sentences) <= n_sentences:
        return sentences
    
    tokens = tokenize(text)
    tf = compute_tf(tokens)
    idf = compute_idf(sentences)
    
    scored = []
    for i, sentence in enumerate(sentences):
        score = score_sentence_tfidf(sentence, tf, idf)
        scored.append((sentence, score, i))
    
    top = sorted(scored, key=lambda x: x[1], reverse=True)[:n_sentences]
    top_sorted = sorted(top, key=lambda x: x[2])
    
    return [s[0] for s in top_sorted]


def categorize_sentences(sentences: List[str]) -> Dict[str, List[str]]:
    """Categorize sentences into thematic groups."""
    goal_keywords = {
        "goal", "objective", "aim", "vision", "purpose", "mission",
        "achieve", "transform", "empower", "enable", "ensure"
    }
    measure_keywords = {
        "implement", "establish", "develop", "provide", "create",
        "deploy", "build", "adopt", "use", "require", "mandate",
        "introduce", "facilitate", "support"
    }
    principle_keywords = {
        "policy", "principle", "framework", "strategy", "approach",
        "standard", "guideline", "compliance", "governance", "regulation"
    }
    
    categories = {"goals": [], "measures": [], "principles": []}
    
    for sentence in sentences:
        lower = sentence.lower()
        g_score = sum(1 for k in goal_keywords if k in lower)
        m_score = sum(1 for k in measure_keywords if k in lower)
        p_score = sum(1 for k in principle_keywords if k in lower)
        
        if g_score >= m_score and g_score >= p_score:
            categories["goals"].append(sentence)
        elif m_score >= p_score:
            categories["measures"].append(sentence)
        else:
            categories["principles"].append(sentence)
    
    all_sentences = sentences.copy()
    if not categories["goals"] and all_sentences:
        categories["goals"] = all_sentences[:2]
    if not categories["measures"] and all_sentences:
        categories["measures"] = all_sentences[2:5]
    if not categories["principles"] and all_sentences:
        categories["principles"] = all_sentences[5:7]
    
    return categories

def build_summary(text: str) -> Dict:
    """Build comprehensive summary from policy document."""
    cleaned = clean_text(text)
    tokens = tokenize(cleaned)
    sentences = sentence_split(cleaned)
    key_sentences = extractive_summarise(cleaned, Config.SUMMARY_SENTENCES)
    categories = categorize_sentences(key_sentences)
    key_terms = extract_key_phrases(cleaned, 20)

    # Build numbered, structured summary text
    summary_lines = []

    summary_lines.append("━━━ MAIN GOALS AND OBJECTIVES ━━━\n")
    goals = categories["goals"][:3]
    for i, sentence in enumerate(goals, 1):
        summary_lines.append(f"  {i}. {sentence.strip()}")

    summary_lines.append("\n━━━ KEY MEASURES AND STRATEGIES ━━━\n")
    measures = categories["measures"][:4]
    for i, sentence in enumerate(measures, 1):
        summary_lines.append(f"  {i}. {sentence.strip()}")

    summary_lines.append("\n━━━ GUIDING PRINCIPLES AND FRAMEWORK ━━━\n")
    principles = categories["principles"][:3]
    for i, sentence in enumerate(principles, 1):
        summary_lines.append(f"  {i}. {sentence.strip()}")

    summary_lines.append("\n━━━ KEY TERMS IDENTIFIED ━━━\n")
    summary_lines.append("  " + "  •  ".join(key_terms[:10]))

    summary_text = "\n".join(summary_lines)

    return {
        "summary": summary_text,
        "categories": {
            "goals": goals,
            "measures": measures,
            "principles": principles
        },
        "key_terms": key_terms,
        "stats": {
            "word_count": len(tokens),
            "sentence_count": len(sentences),
            "key_sentences_extracted": len(key_sentences),
            "compression_ratio": round(
                len(key_sentences) / max(len(sentences), 1) * 100, 1
            )
        }
    }

# GENERATIVE AI - Policy Draft Generation

def generate_with_openai(prompt: str, system_prompt: str) -> str:
    """Generate text using OpenAI's GPT models."""
    if not openai_client:
        raise ValueError("OpenAI client not initialized")
    
    response = openai_client.chat.completions.create(
        model=Config.OPENAI_MODEL,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": prompt}
        ],
        max_tokens=Config.MAX_TOKENS_DRAFT,
        temperature=0.7
    )
    
    return response.choices[0].message.content


def generate_with_anthropic(prompt: str, system_prompt: str) -> str:
    """Generate text using Anthropic's Claude models."""
    if not anthropic_client:
        raise ValueError("Anthropic client not initialized")
    
    response = anthropic_client.messages.create(
        model=Config.ANTHROPIC_MODEL,
        max_tokens=Config.MAX_TOKENS_DRAFT,
        system=system_prompt,
        messages=[
            {"role": "user", "content": prompt}
        ]
    )
    
    return response.content[0].text


# Scenario keywords and adaptations (same as before)
SCENARIO_KEYWORDS = {
    "rural": ["rural", "village", "remote", "kiosk", "sms", "connectivity",
              "low-bandwidth", "offline", "feature phone", "agricultural"],
    "youth": ["youth", "startup", "entrepreneur", "api", "developer", "innovation",
              "young", "tech company", "digital native", "student"],
    "elderly": ["elderly", "older", "disabled", "accessibility", "differently-abled",
                "aged", "senior", "voice", "assisted", "simplified"],
    "crisis": ["crisis", "emergency", "disaster", "relief", "continuity", "outbreak",
               "flood", "conflict", "pandemic", "urgent"],
    "investment": ["investment", "investor", "foreign", "international", "fdi", "trade",
                   "export", "global", "gdpr", "business"],
    "education": ["education", "school", "university", "student", "learning", "teacher",
                  "academic", "curriculum", "training", "skill"],
    "healthcare": ["health", "medical", "hospital", "patient", "doctor", "treatment",
                   "clinic", "medicine", "care", "wellness"]
}

SCENARIO_ADAPTATIONS = {
    "general": {
        "priority_focus": "balanced implementation across all stakeholder groups",
        "objectives": [
            "Ensure equitable access to policy benefits across all demographic groups.",
            "Develop clear implementation guidelines adaptable to local contexts.",
            "Establish monitoring mechanisms to track policy effectiveness.",
            "Create feedback channels for continuous policy improvement.",
            "Align implementation with existing institutional frameworks."
        ],
        "policy_statements": [
            "Implementation shall proceed in phases with regular evaluation checkpoints.",
            "All implementing bodies shall designate a policy coordinator.",
            "Progress reports shall be submitted quarterly to oversight authorities.",
            "Stakeholder consultation shall be conducted before major decisions."
        ],
        "impl_notes": [
            "Phase 1 implementation to focus on quick wins with visible impact.",
            "Resource allocation to be reviewed annually based on implementation progress.",
            "Inter-agency coordination mechanisms to be established within 90 days."
        ]
    },
    "rural": {
        "priority_focus": "offline-first service delivery and community-based access",
        "objectives": [
            "Establish community service kiosks in all rural divisions.",
            "Mandate SMS and USSD-based channels for critical services.",
            "Develop trained community assistants for assisted service delivery.",
            "Prioritise low-bandwidth solutions for 2G network usability.",
            "Allocate dedicated funding for rural infrastructure."
        ],
        "policy_statements": [
            "All new services must include offline fallback mechanisms.",
            "Annual rural inclusion assessments shall be conducted.",
            "Service standards for rural citizens shall match urban outcomes.",
            "Community kiosks shall have official government service status."
        ],
        "impl_notes": [
            "Rollout shall prioritise districts with lowest access indices.",
            "Community assistants shall receive formal accreditation.",
            "Local caching solutions shall be deployed where connectivity is limited."
        ]
    },
    "youth": {
        "priority_focus": "innovation ecosystems, open data, and startup facilitation",
        "objectives": [
            "Establish open APIs for non-sensitive government datasets.",
            "Create sandbox environments for startups to test integrations.",
            "Streamline registration to single online interactions.",
            "Introduce procurement fast-tracks for local solutions.",
            "Align identity standards with international protocols."
        ],
        "policy_statements": [
            "Organisations shall publish API documentation for data services.",
            "Sandbox frameworks shall operate on 12-month cycles.",
            "Youth entrepreneurship shall be a criterion in procurement.",
            "Public registers of data assets shall be maintained quarterly."
        ],
        "impl_notes": [
            "Open data shall comply with data protection legislation.",
            "Fast-tracks shall apply to contracts below defined thresholds.",
            "Annual hackathons shall be funded for solution identification."
        ]
    },
    "elderly": {
        "priority_focus": "accessible, assisted, and inclusive service design",
        "objectives": [
            "Mandate accessibility compliance for all services.",
            "Establish assisted-access counters with trained staff.",
            "Develop voice-interface versions of key services.",
            "Introduce proxy authorisation for family members.",
            "Ensure communications in accessible formats."
        ],
        "policy_statements": [
            "No service shall be exclusively digital.",
            "Accessibility assessments shall be mandatory before deployment.",
            "Information centres shall handle assisted requests.",
            "Accessibility Officers shall be appointed at senior levels."
        ],
        "impl_notes": [
            "Compliance shall be verified by independent audit biennially.",
            "Training shall include communication modules for elderly users.",
            "Proxy frameworks shall be legally binding instruments."
        ]
    },
    "crisis": {
        "priority_focus": "rapid deployment and emergency service continuity",
        "objectives": [
            "Establish pre-authorised emergency data sharing protocols.",
            "Maintain redundant, distributed infrastructure.",
            "Develop and test digital continuity plans.",
            "Create unified emergency service portals.",
            "Mandate regular crisis simulation exercises."
        ],
        "policy_statements": [
            "Data sharing restrictions suspended under emergency status.",
            "Critical services must maintain defined recovery objectives.",
            "Emergency deployments exempt from standard timelines.",
            "Pre-qualified partner rosters shall be maintained."
        ],
        "impl_notes": [
            "Emergency protocols activated by formal declaration.",
            "Post-emergency reviews within 60 days.",
            "Infrastructure shall support automatic failover."
        ]
    },
    "investment": {
        "priority_focus": "international competitiveness and digital trade facilitation",
        "objectives": [
            "Achieve top-quartile international index rankings.",
            "Establish dedicated investor facilitation services.",
            "Publish performance reports to international standards.",
            "Align data protection with international equivalence.",
            "Develop branded credential programmes."
        ],
        "policy_statements": [
            "Business services shall be available in English.",
            "Performance metrics reported using comparable indicators.",
            "Regulatory sandboxes marketed as regional advantages.",
            "Bilateral engagement programmes maintained."
        ],
        "impl_notes": [
            "International alignment coordinated jointly by agencies.",
            "English portals treated as first-class products.",
            "Credential programmes designed with certification bodies."
        ]
    },
    "education": {
        "priority_focus": "digital learning enablement and educational equity",
        "objectives": [
            "Deploy digital learning platforms for all institutions.",
            "Provide teacher training for technology integration.",
            "Establish content repositories with local materials.",
            "Ensure connectivity and device access for students.",
            "Create assessment frameworks for digital competencies."
        ],
        "policy_statements": [
            "Digital literacy integrated into curriculum standards.",
            "Procurement shall prioritise local content compatibility.",
            "Student data privacy protected with enhanced safeguards.",
            "Public-private partnerships encouraged for ed-tech."
        ],
        "impl_notes": [
            "Pilot programmes in diverse educational settings.",
            "Teacher support includes ongoing professional development.",
            "Infrastructure prioritises underserved schools."
        ]
    },
    "healthcare": {
        "priority_focus": "digital health services and patient-centred care",
        "objectives": [
            "Implement electronic health records across providers.",
            "Deploy telemedicine platforms for remote consultation.",
            "Establish health information exchanges.",
            "Develop patient portals for appointments and records.",
            "Create early warning systems for disease surveillance."
        ],
        "policy_statements": [
            "Patient data protected with highest security standards.",
            "Interoperability standards mandated for health IT.",
            "Digital consent mechanisms for data sharing.",
            "Healthcare workers receive digital competency training."
        ],
        "impl_notes": [
            "Implementation complies with health data regulations.",
            "Legacy integration planned with minimal disruption.",
            "Patient feedback informs continuous improvement."
        ]
    }
}


def detect_scenario_type(scenario_text: str) -> str:
    """Detect scenario type from description using keyword matching."""
    lower = scenario_text.lower()
    scores = {}
    
    for scenario_type, keywords in SCENARIO_KEYWORDS.items():
        score = sum(1 for k in keywords if k in lower)
        scores[scenario_type] = score
    
    best = max(scores, key=scores.get)
    return best if scores[best] > 0 else "general"


def generate_with_template(summary: str, scenario_name: str, 
                           scenario_desc: str) -> str:
    """Generate policy draft using template-based approach (fallback)."""
    scenario_type = detect_scenario_type(scenario_desc)
    adaptation = SCENARIO_ADAPTATIONS.get(scenario_type, SCENARIO_ADAPTATIONS["general"])
    
    draft = f"""
╔══════════════════════════════════════════════════════════════════════════════╗
║  ADAPTED POLICY FRAMEWORK                                                    
║  Context: {scenario_name[:60]:<60} ║
╚══════════════════════════════════════════════════════════════════════════════╝

PRIORITY FOCUS: {adaptation['priority_focus'].upper()}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

PREAMBLE

This adapted policy framework addresses the specific requirements of {scenario_name}. 
Building upon the core policy principles, this document adjusts priorities, 
implementation approaches, and resource allocation to meet the distinctive needs 
of the target context while maintaining alignment with overarching policy goals.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

ADAPTED OBJECTIVES

1. {adaptation['objectives'][0]}

2. {adaptation['objectives'][1]}

3. {adaptation['objectives'][2]}

4. {adaptation['objectives'][3]}

5. {adaptation['objectives'][4]}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

KEY POLICY PROVISIONS

• {adaptation['policy_statements'][0]}

• {adaptation['policy_statements'][1]}

• {adaptation['policy_statements'][2]}

• {adaptation['policy_statements'][3]}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

IMPLEMENTATION GUIDELINES

(a) {adaptation['impl_notes'][0]}

(b) {adaptation['impl_notes'][1]}

(c) {adaptation['impl_notes'][2]}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

COMPLIANCE AND REPORTING

This adapted framework operates within the bounds of the source policy document.
Implementation bodies shall report progress quarterly using standard metrics.

[Generated using template-based adaptation | Mode: Offline]
"""
    return draft.strip()


def build_generation_prompt(summary: str, scenario_name: str, 
                            scenario_desc: str) -> Tuple[str, str]:
    """Build prompts for generative AI policy draft creation."""
    system_prompt = """You are an expert policy analyst specialising in adapting 
government policies for different contexts. Generate formal policy documents that:
1. Maintain core policy objectives
2. Adjust priorities for the target scenario
3. Use formal policy language
4. Include specific, actionable provisions
5. Remain realistic and implementable

Format with clear sections: Preamble, Objectives, Policy Provisions, 
Implementation Guidelines, and Compliance requirements."""

    user_prompt = f"""Based on this policy summary, generate an adapted framework for:

SOURCE POLICY SUMMARY:
{summary}

TARGET SCENARIO:
Name: {scenario_name}
Context: {scenario_desc}

Generate a complete adapted policy framework document."""

    return system_prompt, user_prompt


def generate_policy_draft(summary: str, scenario_name: str, 
                          scenario_desc: str) -> Dict:
    """Generate adapted policy draft using configured AI provider."""
    generation_method = "template"
    draft = ""
    
    system_prompt, user_prompt = build_generation_prompt(
        summary, scenario_name, scenario_desc
    )
    
    if Config.AI_PROVIDER == "openai" and openai_client:
        try:
            draft = generate_with_openai(user_prompt, system_prompt)
            generation_method = f"OpenAI ({Config.OPENAI_MODEL})"
        except Exception as e:
            logger.error(f"OpenAI generation failed: {e}")
    
    elif Config.AI_PROVIDER == "anthropic" and anthropic_client:
        try:
            draft = generate_with_anthropic(user_prompt, system_prompt)
            generation_method = f"Anthropic ({Config.ANTHROPIC_MODEL})"
        except Exception as e:
            logger.error(f"Anthropic generation failed: {e}")
    
    if not draft:
        draft = generate_with_template(summary, scenario_name, scenario_desc)
        generation_method = "Template-based (Offline Mode)"
    
    scenario_type = detect_scenario_type(scenario_desc)
    
    return {
        "draft": draft,
        "scenario_name": scenario_name,
        "scenario_type": scenario_type,
        "generation_method": generation_method,
        "ai_provider": Config.AI_PROVIDER
    }


# FLASK ROUTES - API Endpoints

@app.route("/")
def index():
    """Serve the main application page."""
    return send_from_directory("static", "index.html")


@app.route("/api/health", methods=["GET"])
def health_check():
    """Health check endpoint for system status."""
    return jsonify({
        "status": "healthy",
        "ai_provider": Config.AI_PROVIDER,
        "openai_available": openai_client is not None,
        "anthropic_available": anthropic_client is not None,
        "pdf_support": {
            "pdfplumber": PDFPLUMBER_AVAILABLE,
            "pypdf2": PYPDF2_AVAILABLE
        },
        "docx_support": DOCX_AVAILABLE
    })


@app.route("/api/upload", methods=["POST"])
def upload_file():
    """
    Handle file upload and extract text.
    
    Accepts: PDF, DOCX, TXT files
    Returns: Extracted text and metadata
    """
    try:
        if 'file' not in request.files:
            return jsonify({"error": "No file provided"}), 400
        
        file = request.files['file']
        
        if file.filename == '':
            return jsonify({"error": "No file selected"}), 400
        
        if not allowed_file(file.filename):
            return jsonify({
                "error": f"File type not supported. Allowed: {', '.join(ALLOWED_EXTENSIONS)}"
            }), 400
        
        # Process the file
        text, metadata = process_uploaded_file(file)
        
        if not text or len(text.strip()) < 100:
            return jsonify({
                "error": "Could not extract sufficient text from the file. The document may be image-based or empty."
            }), 400
        
        logger.info(f"Successfully extracted {len(text)} characters from {metadata['filename']}")
        
        return jsonify({
            "text": text,
            "metadata": metadata,
            "success": True
        })
    
    except Exception as e:
        logger.error(f"File upload error: {e}")
        return jsonify({"error": f"Error processing file: {str(e)}"}), 500


@app.route("/api/summarise", methods=["POST"])
def summarise():
    """Summarise policy document text."""
    try:
        data = request.get_json(force=True)
        text = (data.get("text") or "").strip()
        
        if not text:
            return jsonify({"error": "No policy text provided."}), 400
        
        if len(text) < 100:
            return jsonify({
                "error": "Text too short. Please provide a more substantial policy document."
            }), 400
        
        result = build_summary(text)
        logger.info(f"Generated summary: {result['stats']['word_count']} words processed")
        
        return jsonify(result)
    
    except Exception as e:
        logger.error(f"Summarisation error: {e}")
        return jsonify({"error": f"Processing error: {str(e)}"}), 500


@app.route("/api/generate", methods=["POST"])
def generate():
    """Generate scenario-based policy draft."""
    try:
        data = request.get_json(force=True)
        summary = (data.get("summary") or "").strip()
        scenario_name = (data.get("scenario_name") or "Custom Scenario").strip()
        scenario_desc = (data.get("scenario") or "").strip()
        
        if not summary:
            return jsonify({"error": "Summary is required."}), 400
        
        if not scenario_desc:
            return jsonify({"error": "Scenario description is required."}), 400
        
        result = generate_policy_draft(summary, scenario_name, scenario_desc)
        logger.info(f"Generated policy draft for: {scenario_name}")
        
        return jsonify(result)
    
    except Exception as e:
        logger.error(f"Generation error: {e}")
        return jsonify({"error": f"Generation error: {str(e)}"}), 500


@app.route("/api/scenarios", methods=["GET"])
def get_scenarios():
    """Get available preset scenarios."""
    presets = [
        {"id": "rural", "name": "Rural Digital Inclusion",
         "description": "Adapt for rural areas with limited connectivity."},
        {"id": "youth", "name": "Youth and Startups",
         "description": "Reframe for young entrepreneurs and startups."},
        {"id": "elderly", "name": "Elderly and Differently-Abled",
         "description": "Adapt for accessibility and assisted services."},
        {"id": "crisis", "name": "Crisis and Emergency Response",
         "description": "Adapt for emergency conditions."},
        {"id": "investment", "name": "International Investment",
         "description": "Reposition for foreign investment attraction."},
        {"id": "education", "name": "Educational Transformation",
         "description": "Adapt for educational institutions."},
        {"id": "healthcare", "name": "Healthcare Digital Services",
         "description": "Adapt for healthcare sector."}
    ]
    return jsonify({"scenarios": presets})


# MAIN ENTRY POINT

if __name__ == "__main__":
    print("\n" + "="*70)
    print("PolicyLens - Policy Summarisation & Scenario Generation")
    print("="*70)
    print(f"AI Provider: {Config.AI_PROVIDER}")
    print(f"PDF Support: pdfplumber={PDFPLUMBER_AVAILABLE}, PyPDF2={PYPDF2_AVAILABLE}")
    print(f"DOCX Support: {DOCX_AVAILABLE}")
    print(f"Running on: http://localhost:{Config.PORT}")
    print("="*70 + "\n")
    
    app.run(debug=Config.DEBUG, port=Config.PORT)